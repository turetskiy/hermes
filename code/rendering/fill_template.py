"""
Generic filler: data/template.docx + a content dict -> a final, track/vacancy-specific docx.
Opens the template fresh each run (never mutates it). Called by assemble.py / pipeline.py.

Responsibilities split several ways:
  docx_write.py   - writing text/fonts/links into a paragraph (+ the per-style meta sidecar)
  docx_expand.py  - cloning one anchor paragraph into N (skills clusters, role bullets)
  docx_fixed.py   - filling NAME/CONTACT/COMPANY/EDU_* tokens from identity.json, doc-wide
  docx_layout.py  - finding/removing placeholder paragraphs and post-fill whitespace tidy
  this file       - orchestration: roles, list sections, and build()

Content schema:
  identity: {name, contact, companies:{role1..4}, education:{degree,institution,dates}} (optional -
    a token with no matching value is left as-is; see rendering/docx_fixed.py)
  tagline: str | summary: str (**bold** ok) | skills: [{label, items}, ...]  (flat list of clusters)
  roles: {role1..role4: {title, dates, bullets:[...]} | null} | public_speaking:[...] | articles:[...]
  footer_title: str
Empty public_speaking/articles list -> that whole section is removed; a role with no bullets is
dropped entirely. A role's bullets have no capacity ceiling - the template needs only ONE {{ROLEn_B1}}
anchor paragraph per role, expanded to however many bullets are given (see rendering/docx_expand.py);
how many that actually is is a content/track decision (positioning.json's per-role max_bullets), not
a template one.

CLI:  python -m rendering.fill_template <content.json> <output.docx>      (run from code/)
"""
import json
import os
import sys

if __package__ in (None, ""):  # allow `python code/rendering/fill_template.py ...` direct runs
    sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import docx
from docx.shared import Pt

import paths
from rendering import docx_write as W
from rendering import docx_expand as E
from rendering import docx_fixed as F
from rendering import docx_layout as L

TEMPLATE = os.path.join(paths.DATA, "template.docx")
SPEAK_CAPACITY = 4
ARTICLE_CAPACITY = 2


def _clusters(skills):
    """Skills as a flat list of {label, items} clusters, for the single {{SKILLS}} anchor. Accepts a
    list (the current shape) or a legacy dict-of-groups, so any number of skill groups just works."""
    if isinstance(skills, list):
        return skills
    if isinstance(skills, dict):
        return [c for group in skills.values() if isinstance(group, list) for c in group]
    return []


def fill_role(doc, role_key, role_content):
    title_p = L.find_role_title_paragraph(doc, role_key)
    bullet_anchor = L.find_placeholder_paragraph(doc, f"{{{{{role_key.upper()}_B1}}}}")

    if role_content is None:
        L.remove_paragraph(title_p)
        L.remove_paragraph(bullet_anchor)
        return

    title_token = f"{{{{{role_key.upper()}_TITLE}}}}"
    dates_token = f"{{{{{role_key.upper()}_DATES}}}}"
    replaced_title = replaced_dates = False
    for r in title_p.runs:
        if title_token in r.text:
            r.text = r.text.replace(title_token, role_content["title"])
            replaced_title = True
        elif dates_token in r.text:
            r.text = r.text.replace(dates_token, role_content["dates"])
            replaced_dates = True
    if not (replaced_title and replaced_dates):
        raise ValueError(f"{role_key}: title/dates placeholder run not found (title={replaced_title}, dates={replaced_dates})")

    E.expand_lines(bullet_anchor, role_content["bullets"], W.set_segments)


def fill_list_section(doc, prefix, capacity, items, section_header_text):
    placeholders = [L.find_placeholder_paragraph(doc, f"{{{{{prefix}_{i}}}}}") for i in range(1, capacity + 1)]
    if len(items) > capacity:
        raise ValueError(f"{prefix}: {len(items)} items given, template only has {capacity} slots")
    if not items:
        header_p = next(p for p in doc.paragraphs if p.text.strip() == section_header_text)
        L.remove_paragraph(header_p)
        for ph in placeholders:
            L.remove_paragraph(ph)
        return
    for i, ph in enumerate(placeholders):
        if i < len(items):
            W.set_list_item(ph, items[i])
        else:
            L.remove_paragraph(ph)


def build(content: dict, output_path: str):
    W.load_meta()
    doc = docx.Document(TEMPLATE)
    F.fill_fixed_tokens(doc, content.get("identity", {}))

    tagline_p = L.find_placeholder_paragraph(doc, "{{TAGLINE}}")
    W.set_segments(tagline_p, content["tagline"])
    W.apply_tagline_style(tagline_p)
    W.set_segments(L.find_placeholder_paragraph(doc, "{{SUMMARY}}"), content["summary"])
    E.expand_labeled_lines(L.find_placeholder_paragraph(doc, "{{SKILLS}}"), _clusters(content.get("skills")))

    for role_key in ("role1", "role2", "role3", "role4"):
        fill_role(doc, role_key, content["roles"].get(role_key))

    fill_list_section(doc, "SPEAK", SPEAK_CAPACITY, content.get("public_speaking", []),
                      W.META.get("speak_header", "Public speeches"))
    fill_list_section(doc, "ARTICLE", ARTICLE_CAPACITY, content.get("articles", []),
                      W.META.get("articles_header", "Articles"))

    footer_p = doc.sections[0].footer.paragraphs[0]
    W.set_segments(footer_p, content["footer_title"])
    for r in footer_p.runs:
        r.font.size = Pt(W.META.get("footer_size", 8))

    L.collapse_blank_runs(doc)
    L.add_section_spacing(doc)
    L.normalize_line_spacing(doc)
    doc.save(output_path)
    print("Saved:", output_path)


if __name__ == "__main__":
    content_path, output_path = sys.argv[1], sys.argv[2]
    with open(content_path) as f:
        content = json.load(f)
    build(content, output_path)
