"""
docx_layout.py - document-structure operations on the filled resume: finding placeholder
paragraphs, removing them, and the post-fill whitespace tidy (uniform line spacing, section
gaps, collapsing double blanks). No text/font writing (that is docx_write.py); no content
orchestration (that is fill_template.py) - just paragraph plumbing.
"""
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def find_role_title_paragraph(doc, role_key):
    token = f"{{{{{role_key.upper()}_TITLE}}}}"
    for p in doc.paragraphs:
        if p.text.startswith(token):
            return p
    raise ValueError(f"role title placeholder {token} not found")


def find_placeholder_paragraph(doc, token):
    for p in doc.paragraphs:
        if p.text.strip() == token:
            return p
    raise ValueError(f"placeholder {token} not found (template may have been hand-edited)")


def remove_paragraph(p):
    p._element.getparent().remove(p._element)


def normalize_line_spacing(doc):
    """Uniform SINGLE line spacing on every body/footer line."""
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for p in doc.sections[0].footer.paragraphs:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_section_spacing(doc):
    """Give any section header or role subheader that butts directly against non-blank content
    above it the same gap the other sections have: a blank paragraph before it. (space_before is
    avoided - on a header it shifts the text but leaves the underline rule behind.) Fills only gaps."""
    paras = doc.paragraphs
    targets = []
    for i, p in enumerate(paras):
        if i == 0:
            continue
        prev = paras[i - 1]
        if prev.text.strip() == "":
            continue  # already has a gap
        is_header = p.style.name == "CvExpHeader"
        has_link = p._p.find(qn("w:hyperlink")) is not None
        is_role = p.style.name == "Normal" and has_link and "\t" in p.text
        if is_role and prev.style.name == "CvExpHeader":
            continue  # first role right under its "Experience" header - no gap wanted
        if is_header or is_role:
            targets.append(p)
    for p in targets:
        p._p.addprevious(OxmlElement("w:p"))  # insert a blank Normal paragraph before it


def collapse_blank_runs(doc):
    """After deletions, collapse any run of 2+ consecutive empty paragraphs to 1."""
    to_remove = []
    prev_blank = False
    for p in doc.paragraphs:
        is_blank = p.text.strip() == ""
        if is_blank and prev_blank:
            to_remove.append(p)
        prev_blank = is_blank
    for p in to_remove:
        remove_paragraph(p)
