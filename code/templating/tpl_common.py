"""
tpl_common.py - shared docx helpers for the template builders (build_template.py + tpl_skeleton.py).

Low-level, style-agnostic: page setup, the Normal + CvExpHeader styles, run/paragraph builders,
a bottom-border rule, a right tab stop, and character spacing. All visual identity lives in the
THEME dicts in build_template.py; this module only knows how to render what a theme asks for.
"""
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

A4_W, A4_H = 21.0, 29.7          # cm - A4 (python-docx defaults to Letter, so set it explicitly)
INK = "222529"
MUTED = "6B727C"
HEADER_STYLE = "CvExpHeader"     # the style name fill_template/assemble use to detect section headers


def setup_page(doc, margins):
    top, bottom, left, right = margins
    for s in doc.sections:
        s.page_width, s.page_height = Cm(A4_W), Cm(A4_H)
        s.top_margin, s.bottom_margin = Cm(top), Cm(bottom)
        s.left_margin, s.right_margin = Cm(left), Cm(right)


def normal_style(doc, font, size):
    st = doc.styles["Normal"]
    st.font.name = font
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(INK)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def ensure_header_style(doc):
    if HEADER_STYLE not in [s.name for s in doc.styles]:
        st = doc.styles.add_style(HEADER_STYLE, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = doc.styles["Normal"]


def char_spacing(r, twips):
    rPr = r._r.get_or_add_rPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:val"), str(int(twips)))
    rPr.append(sp)


def run(p, text, font=None, size=None, bold=False, italic=False, color=None, caps=False, spacing=None):
    r = p.add_run(text)
    if font:
        r.font.name = font
    if size:
        r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if caps:
        r.font.all_caps = True
    if spacing is not None:
        char_spacing(r, spacing)
    return r


def para(doc, space_before=0, space_after=0, style=None, align=None):
    p = doc.add_paragraph()
    if style:
        p.style = style
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align is not None:
        pf.alignment = align
    return p


def blank(doc):
    return doc.add_paragraph()


def bottom_border(p, color, sz=6, space=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(sz))
    b.set(qn("w:space"), str(space))
    b.set(qn("w:color"), color)
    pbdr.append(b)
    pPr.append(pbdr)


def right_tab(p, pos_cm):
    p.paragraph_format.tab_stops.add_tab_stop(Cm(pos_cm), WD_TAB_ALIGNMENT.RIGHT)
