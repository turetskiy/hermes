"""Extract plain text from a single file. docx/pdf by extension; everything else is read as text if it
actually looks like text - so a note with no extension (or an odd one) is NOT silently skipped. Truly
binary files (images, zip, legacy .doc, ...) return None. Text/docx need no extra deps; PDF needs pypdf."""
import os

TEXT_EXT = {".txt", ".md", ".markdown", ".rst", ".json", ".csv", ".tsv", ".log", ".yaml", ".yml"}


def extract(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return _docx(path)
    if ext == ".pdf":
        return _pdf(path)
    if ext in TEXT_EXT or _looks_text(path):
        return _text(path)
    return None  # binary / unreadable -> caller records it as skipped


def _looks_text(path):
    try:
        with open(path, "rb") as f:
            chunk = f.read(4096)
    except OSError:
        return False
    if b"\x00" in chunk:  # a NUL byte means binary
        return False
    try:
        chunk.decode("utf-8")
        return True
    except UnicodeDecodeError:
        return False


def _text(path):
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


def _docx(path):
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _pdf(path):
    try:
        from pypdf import PdfReader
    except ImportError:
        raise RuntimeError("PDF support needs pypdf - run: pip install pypdf")
    reader = PdfReader(path)
    return "\n".join((page.extract_text() or "") for page in reader.pages)
